#!/usr/bin/env python3
"""
render_mrc_handbook.py

Converts each family of generated Maintenance Requirement Cards (MRC) --
execution-plan/mrc-cards/{master,network-infra,ops}/MRC*.md -- into one
clean, formatted DOCX handbook per family, with real Word tables (not
markdown pipe-tables), a clickable table of contents, a page break between
cards, and professional typography (design-foundations Nexus palette,
Calibri throughout -- a system font, since DOCX must render correctly on
the viewer's machine without embedded fonts).

THIS SCRIPT IS A PURE RENDERER. It does not compute, resolve, or invent
any content -- it reads the already-generated .md cards verbatim and
reflows their existing markdown (headings, tables, bullet/numbered lists,
bold text) into DOCX equivalents. If a card's *content* is wrong, fix the
card's generator (build_mrc_cards.py / build_operational_tasking.py /
build_network_infra_tasking.py) and regenerate the .md cards first, then
re-run this script -- do not hand-edit the .docx output, and do not
hand-edit this script's per-card content logic without a corresponding
source-of-truth change upstream.

Regenerate any time with:
    python3 execution-plan/tools/render_mrc_handbook.py

Produces:
    execution-plan/mrc-cards/master/MRC-HANDBOOK-MASTER.docx
    execution-plan/mrc-cards/network-infra/MRC-HANDBOOK-NETWORK-INFRA.docx
    execution-plan/mrc-cards/ops/MRC-HANDBOOK-OPS.docx

These .docx files ARE committed to source control as the distributable
handbook deliverable -- see .gitignore's comment for the rationale. Only the
.pdf companion (produced separately via `soffice --headless --convert-to
pdf`) is gitignored, since it is purely a rendered-from-the-.docx artifact.
After regenerating, re-commit the refreshed .docx -- don't leave a stale
copy in source control.
"""
import os
import re
import sys
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

TOOLS_DIR = os.path.dirname(os.path.abspath(__file__))
EXEC_PLAN_DIR = os.path.dirname(TOOLS_DIR)
CARDS_ROOT = os.path.join(EXEC_PLAN_DIR, "mrc-cards")

# ---------------------------------------------------------------------------
# Nexus palette (light mode) -- see design-foundations skill
# ---------------------------------------------------------------------------
COLOR_TEXT = RGBColor(0x28, 0x25, 0x1D)
COLOR_TEXT_MUTED = RGBColor(0x7A, 0x79, 0x74)
COLOR_PRIMARY = RGBColor(0x01, 0x69, 0x6F)
COLOR_PRIMARY_HOVER = RGBColor(0x0C, 0x4E, 0x54)
COLOR_BORDER = "D4D1CA"
COLOR_SURFACE = "F9F8F5"
COLOR_HEADER_FILL = "01696F"

FAMILIES = [
    {
        "key": "master",
        "dir": os.path.join(CARDS_ROOT, "master"),
        "prefix": "MRC-",
        "index": "INDEX.md",
        "title": "JSIG Master Calendar",
        "subtitle": "Maintenance Requirement Card Handbook \u2014 110 JSIG-Driven Tasks",
        "out": "MRC-HANDBOOK-MASTER.docx",
        "source_note": "Generated from execution-plan/mrc-cards/master/*.md, produced by "
                        "execution-plan/tools/build_mrc_cards.py from MAINTENANCE-PLAN.md \u00a74.",
    },
    {
        "key": "network-infra",
        "dir": os.path.join(CARDS_ROOT, "network-infra"),
        "prefix": "MRC-NET-",
        "index": "INDEX.md",
        "title": "Network Infrastructure Tasking",
        "subtitle": "Maintenance Requirement Card Handbook \u2014 16 Switch/Router/Firewall Health Tasks",
        "out": "MRC-HANDBOOK-NETWORK-INFRA.docx",
        "source_note": "Generated from execution-plan/mrc-cards/network-infra/*.md, produced by "
                        "execution-plan/tools/build_network_infra_tasking.py from "
                        "NETWORK-INFRASTRUCTURE-TASKING.md.",
    },
    {
        "key": "ops",
        "dir": os.path.join(CARDS_ROOT, "ops"),
        "prefix": "MRC-OPS-",
        "index": "INDEX.md",
        "title": "Operational Tasking",
        "subtitle": "Maintenance Requirement Card Handbook \u2014 34 Non-JSIG Functional/Health Tasks",
        "out": "MRC-HANDBOOK-OPS.docx",
        "source_note": "Generated from execution-plan/mrc-cards/ops/*.md, produced by "
                        "execution-plan/tools/build_operational_tasking.py from "
                        "OPERATIONAL-TASKING.md.",
    },
]

# ---------------------------------------------------------------------------
# Inline markdown -> runs (bold **x**, code `x`, links [x](y) -> plain text x)
# ---------------------------------------------------------------------------
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")


def add_inline_runs(paragraph, text, base_size=10.5, color=None, muted_ok=True):
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
        elif chunk.startswith("[") and "](" in chunk:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            label = m.group(1) if m else chunk
            run = paragraph.add_run(label)
            run.italic = True
            if muted_ok:
                run.font.color.rgb = COLOR_TEXT_MUTED
        else:
            run = paragraph.add_run(chunk)
        run.font.size = Pt(base_size)
        run.font.name = "Calibri"
        if color is not None and not (chunk.startswith("[") and "](" in chunk):
            run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Card markdown parser
# ---------------------------------------------------------------------------
def split_sections(md_text):
    """Split a card's markdown into (title_line, blockquote, {section_name: [lines]})."""
    lines = md_text.splitlines()
    title = lines[0].lstrip("# ").strip() if lines and lines[0].startswith("# ") else ""
    idx = 1
    blockquote = []
    while idx < len(lines) and (lines[idx].startswith(">") or not lines[idx].strip()):
        if lines[idx].startswith(">"):
            blockquote.append(lines[idx].lstrip("> ").strip())
        idx += 1
    sections = {}
    current = None
    for line in lines[idx:]:
        m = re.match(r"^##\s+\d+\.\s*(.+)$", line)
        if m:
            current = m.group(1).strip()
            sections[current] = []
        elif current is not None:
            sections[current].append(line)
    return title, " ".join(blockquote), sections


def parse_table(block_lines):
    """Parse a contiguous list of markdown '|'-table lines into (header, rows)."""
    rows = []
    for line in block_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s:-]+\|$", line) or set(line.replace("|", "").strip()) <= set("-: "):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return None, []
    return rows[0], rows[1:]


def render_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        shade_cell(hdr_cells[i], COLOR_HEADER_FILL)
        p = hdr_cells[i].paragraphs[0]
        p.text = ""
        run = p.add_run(re.sub(r"[*`]", "", h))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    for r in rows:
        cells = table.add_row().cells
        for i in range(len(header)):
            val = r[i] if i < len(r) else ""
            p = cells[i].paragraphs[0]
            add_inline_runs(p, val, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), COLOR_BORDER)
        borders.append(el)
    tblPr.append(borders)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_bookmark(paragraph, name, bmk_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bmk_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bmk_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "01696F")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")
    rpr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rpr.append(rFonts)
    run_el.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def render_body_lines(doc, lines):
    """Render freeform section body: paragraphs, bullet lists, numbered lists, tables."""
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            header, rows = parse_table(block)
            if header:
                render_table(doc, header, rows)
            continue
        if re.match(r"^-\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, re.sub(r"^-\s+", "", stripped))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            add_inline_runs(p, stripped, base_size=11, color=COLOR_PRIMARY_HOVER)
            p.paragraph_format.space_before = Pt(6)
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        p.paragraph_format.space_after = Pt(6)
        i += 1


def set_base_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_TEXT
    style.paragraph_format.space_after = Pt(4)
    for name in ("List Bullet", "List Number"):
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(10.5)
        s.font.color.rgb = COLOR_TEXT


def add_heading(doc, text, level=1, size=16, color=None, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = color or COLOR_PRIMARY
    return p


def build_handbook(family):
    card_dir = family["dir"]
    prefix = family["prefix"]
    files = sorted(
        f for f in os.listdir(card_dir)
        if f.startswith(prefix) and f.endswith(".md")
    )
    if not files:
        print(f"  no cards found in {card_dir}, skipping")
        return None

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    set_base_styles(doc)

    # ---- Cover page ----
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Maintenance Requirement Card Handbook")
    run.font.size = Pt(15)
    run.font.color.rgb = COLOR_TEXT_MUTED
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(family["title"])
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(family["subtitle"])
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_TEXT_MUTED
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run(f"Generated {date.today().isoformat()} \u00b7 {len(files)} cards")
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEXT_MUTED
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(family["source_note"])
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = COLOR_TEXT_MUTED
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "This handbook is a derived, regenerable artifact rendered from the source .md cards. "
        "Do not hand-edit -- edit the calendar/generator and re-run "
        "execution-plan/tools/render_mrc_handbook.py."
    )
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = COLOR_TEXT_MUTED
    run.font.name = "Calibri"

    doc.add_page_break()

    # ---- Table of contents ----
    add_heading(doc, "Table of Contents", size=20, space_before=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    parsed_cards = []
    for fname in files:
        path = os.path.join(card_dir, fname)
        with open(path, encoding="utf-8") as f:
            md_text = f.read()
        title, blockquote, sections = split_sections(md_text)
        parsed_cards.append((fname, title, blockquote, sections))

    for i, (fname, title, _bq, _sections) in enumerate(parsed_cards):
        anchor = f"card_{i}"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_internal_hyperlink(p, title, anchor)

    doc.add_page_break()

    # ---- Cards ----
    for i, (fname, title, blockquote, sections) in enumerate(parsed_cards):
        anchor = f"card_{i}"
        heading_p = add_heading(doc, title, size=17, space_before=0, space_after=10)
        add_bookmark(heading_p, anchor, i)

        if blockquote:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(blockquote)
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = COLOR_TEXT_MUTED
            run.font.name = "Calibri"

        for section_name, body_lines in sections.items():
            add_heading(doc, section_name, size=12.5, color=COLOR_PRIMARY_HOVER,
                        space_before=14, space_after=4)
            render_body_lines(doc, body_lines)

        if i < len(parsed_cards) - 1:
            doc.add_page_break()

    out_path = os.path.join(card_dir, family["out"])
    doc.save(out_path)
    return out_path, len(files)


def main():
    print(f"Rendering MRC handbooks from {CARDS_ROOT}")
    for family in FAMILIES:
        print(f"- {family['key']} ...")
        result = build_handbook(family)
        if result:
            out_path, count = result
            print(f"  wrote {out_path} ({count} cards)")


if __name__ == "__main__":
    main()
